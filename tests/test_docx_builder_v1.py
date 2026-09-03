"""
Test DOCX Builder v1 — Verifica la generación nativa de documentos .docx desde SessionDocumentV1.
Ejecutar: python tests/test_docx_builder_v1.py
"""
import os
import sys
import json
import io
import zipfile
from xml.etree import ElementTree as ET

# Force UTF-8 on Windows
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

TEST_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(TEST_DIR)
sys.path.insert(0, os.path.join(ROOT_DIR, 'backend'))

from docx import Document
from models.session_document import SessionDocumentV1
from docx_builder_v1 import build_docx_from_v1
from docx.oxml.ns import qn


def test_fixture_docx(fixture_filename: str):
    """Carga un fixture v1 y genera un documento DOCX válido."""
    print(f"\nGenerando DOCX para: {fixture_filename}...")
    fixture_path = os.path.join(TEST_DIR, 'fixtures', fixture_filename)

    with open(fixture_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc_v1 = SessionDocumentV1(**data)
    docx_stream = build_docx_from_v1(doc_v1)

    assert isinstance(docx_stream, io.BytesIO), "Debe devolver io.BytesIO"
    size = docx_stream.getbuffer().nbytes
    assert size > 5000, f"El archivo DOCX debe tener al menos 5KB (generado: {size} bytes)"

    # Validar que python-docx puede reabrir el stream sin errores
    docx_stream.seek(0)
    reloaded_doc = Document(docx_stream)
    assert len(reloaded_doc.tables) >= 5, f"Debe tener al menos 5 tablas principales (tiene {len(reloaded_doc.tables)})"
    assert any(p.text.startswith("SESIÓN DE APRENDIZAJE N°") for p in reloaded_doc.paragraphs), (
        "Debe incluir el encabezado oficial de sesión antes de las tablas"
    )

    print(f"  ✓ {fixture_filename}: DOCX generado exitosamente ({size} bytes, {len(reloaded_doc.tables)} tablas)")
    return size


def test_custom_presentation_tokens():
    fixture_path = os.path.join(TEST_DIR, 'fixtures', 'secundaria-matematica-polya.v1.json')
    with open(fixture_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    data['presentation'] = {
        'preset': 'moderno', 'primaryColor': '#334155', 'accentColor': '#0F766E',
        'headerBackground': '#CCFBF1', 'fontFamily': 'Calibri', 'fontSizePt': 11,
        'cellPadding': 'comfortable', 'lineHeight': 1.3,
    }
    output = build_docx_from_v1(SessionDocumentV1(**data))
    rendered = Document(output)

    fills = {
        node.get(qn('w:fill'))
        for table in rendered.tables
        for node in table._tbl.xpath('.//w:shd')
    }
    borders = {
        node.get(qn('w:color'))
        for table in rendered.tables
        for node in table._tbl.xpath('.//w:tblBorders/*')
    }
    fonts = {
        run.font.name
        for table in rendered.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    }
    assert 'CCFBF1' in fills, 'El fondo de cabecera personalizado debe llegar a Word'
    assert borders == {'334155'}, f'Los bordes deben usar el color primario: {borders}'
    assert fonts == {'Calibri'}, f'Todas las fuentes deben usar Calibri: {fonts}'


def test_template_fidelity_and_content_coverage():
    fixture_path = os.path.join(TEST_DIR, 'fixtures', 'secundaria-matematica-polya.v1.json')
    with open(fixture_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    data['listaCotejo']['alumnos'] = []

    output = build_docx_from_v1(SessionDocumentV1(**data))
    raw = output.getvalue()
    rendered = Document(io.BytesIO(raw))
    all_text = ' '.join(
        node.text or ''
        for node in rendered._element.body.iter()
        if node.tag == qn('w:t')
    )

    assert abs(rendered.sections[0].page_width.inches - 8.27) < 0.01
    assert abs(rendered.sections[0].page_height.inches - 11.69) < 0.01
    assert rendered.sections[0].header._element.xpath('.//w:drawing | .//w:pict'), (
        'Debe conservar el encabezado institucional de la plantilla'
    )
    for expected in (
        data['metadata']['dre'], data['metadata']['ugel'], data['metadata']['director'],
        data['proposito']['desempeno'], data['evaluacion']['criterioConsolidado'],
        data['momentos']['desarrollo']['procesos'][0]['titulo']
    ):
        assert expected.casefold() in all_text.casefold(), f'Contenido v1 perdido en DOCX: {expected!r}'

    checklist = next(table for table in rendered.tables if len(table.columns) == 10)
    assert len(checklist.rows) == 32, 'Debe crear 30 filas vacías aunque no haya nómina'
    assert checklist.rows[0]._tr.xpath('./w:trPr/w:tblHeader')
    assert checklist.rows[1]._tr.xpath('./w:trPr/w:tblHeader')

    with zipfile.ZipFile(io.BytesIO(raw)) as package:
        xml_parts = b'\n'.join(
            package.read(name) for name in package.namelist() if name.endswith('.xml')
        )
        for private_sample in (b'Edward Lenin', b'ALAYA TRIGOSO', b'TRANSFORMACIONES GEOM'):
            assert private_sample not in xml_parts, 'La plantilla sanitizada filtró datos de ejemplo'

        styles_root = ET.fromstring(package.read('word/styles.xml'))
        document_root = ET.fromstring(package.read('word/document.xml'))
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        declared = {
            node.attrib[qn('w:styleId')]
            for node in styles_root.findall('w:style', ns)
        }
        referenced = {
            node.attrib[qn('w:val')]
            for selector in ('.//w:pStyle', './/w:rStyle', './/w:tblStyle')
            for node in document_root.findall(selector, ns)
            if qn('w:val') in node.attrib
        }
        assert referenced <= declared, f'Estilos OOXML faltantes: {sorted(referenced - declared)}'


if __name__ == '__main__':
    print("=" * 60)
    print("TEST DOCX BUILDER V1 — SessionDocument v1 → .docx")
    print("=" * 60)

    size1 = test_fixture_docx('secundaria-matematica-polya.v1.json')
    size2 = test_fixture_docx('inicial.v1.json')
    size3 = test_fixture_docx('primaria-con-ficha.v1.json')
    test_custom_presentation_tokens()
    test_template_fidelity_and_content_coverage()

    print("\n" + "=" * 60)
    print(">>> TODOS LOS TESTS DE DOCX BUILDER V1 PASARON EXITOSAMENTE <<<")
    print("=" * 60)
