"""Pruebas del render híbrido LaTeX → OMML → imagen → texto."""
from __future__ import annotations

import io
import json
import sys
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from lxml import etree
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
sys.path.insert(0, str(BACKEND))

from models.session_document import SessionDocumentV1  # noqa: E402
from docx_builder import append_html_to_cell_or_paragraph  # noqa: E402
from docx_builder_v1 import build_docx_from_v1  # noqa: E402
from word_math import append_mixed_text, latex_to_omml, split_latex_segments  # noqa: E402


def _document_xml(stream: io.BytesIO) -> str:
    stream.seek(0)
    with zipfile.ZipFile(stream) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_segment_detection():
    segments = list(split_latex_segments(r"Calcula $x^2$ y luego $$\frac{1}{2}$$."))
    formulas = [(value, display) for is_math, value, display, _ in segments if is_math]
    assert formulas == [("x^2", False), (r"\frac{1}{2}", True)]


def test_native_omml_structures():
    equation = latex_to_omml(r"x_{1}^{2}+\frac{3}{4}+\sqrt{y}")
    xml = etree.tostring(equation, encoding="unicode")
    assert "m:sSubSup" in xml
    assert "m:f" in xml
    assert "m:rad" in xml


def test_html_keeps_formatting_and_adds_native_equation():
    doc = Document()
    paragraph = doc.add_paragraph()
    append_html_to_cell_or_paragraph(
        paragraph,
        r"<p><strong>Resolvemos:</strong> $2x+5=21 \Rightarrow x=8$</p>",
    )
    stream = io.BytesIO()
    doc.save(stream)
    xml = _document_xml(stream)
    assert "<m:oMath" in xml
    assert "<w:b" in xml
    assert r"\Rightarrow" not in xml


def test_unknown_native_command_uses_image():
    doc = Document()
    routes = append_mixed_text(doc.add_paragraph(), r"Valor: $\sin(x)+\cos(x)$")
    assert "image" in routes
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    with zipfile.ZipFile(stream) as archive:
        media_name = next(name for name in archive.namelist() if name.startswith("word/media/"))
        with Image.open(io.BytesIO(archive.read(media_name))) as image:
            assert image.mode == "RGBA"
            assert image.getchannel("A").getextrema()[0] == 0


def test_last_resort_preserves_original_latex():
    doc = Document()
    with patch("word_math._render_formula_png", return_value=None):
        routes = append_mixed_text(doc.add_paragraph(), r"$\comandoInventado{x}$")
    assert routes == ["latex-text"]
    assert doc.paragraphs[0].text == r"$\comandoInventado{x}$"


def test_v1_builder_integration():
    fixture = json.loads(
        (ROOT / "tests" / "fixtures" / "secundaria-matematica-polya.v1.json")
        .read_text(encoding="utf-8")
    )
    fixture["momentos"]["desarrollo"]["procesos"][0]["contenido"]["value"] = (
        r"<p>Aplicamos $x^2+\frac{1}{2}$ y comparamos con $\sin(\theta)$.</p>"
    )
    stream = build_docx_from_v1(SessionDocumentV1(**fixture))
    xml = _document_xml(stream)
    assert "<m:oMath" in xml
    stream.seek(0)
    with zipfile.ZipFile(stream) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())
    stream.seek(0)
    Document(stream)  # El paquete sigue siendo un DOCX válido.


if __name__ == "__main__":
    tests = [
        test_segment_detection,
        test_native_omml_structures,
        test_html_keeps_formatting_and_adds_native_equation,
        test_unknown_native_command_uses_image,
        test_last_resort_preserves_original_latex,
        test_v1_builder_integration,
    ]
    for test in tests:
        test()
        print(f"✓ {test.__name__}")
    print("test_word_math.py: OK")
