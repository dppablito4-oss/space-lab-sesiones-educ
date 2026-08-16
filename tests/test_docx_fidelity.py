import sys
import os
import io
import docx
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))
from docx_builder import build_docx_from_json
from main import normalize_sesion_data, SesionAprendizajeRequest

raw_payload = {
    "metadata": {
        "institucion": "I.E. 3082 SEÑOR DE LOS MILAGROS",
        "dre": "LIMA METROPOLITANA",
        "ugel": "UGEL 02",
        "docente": "LIC. PABLO PÉREZ",
        "director": "MG. CARLOS MENDOZA",
        "fecha": "16/08/2026",
        "nivel": "SECUNDARIA",
        "numero_sesion": "04",
        "grado": "4°",
        "seccion": "A",
        "area": "MATEMÁTICA",
        "duracion": "90 min",
        "unidad": "UNIDAD 01",
        "titulo": "RESOLVEMOS PROBLEMAS DE REGULARIDAD, EQUIVALENCIA Y CAMBIO CON SISTEMAS DE ECUACIONES"
    },
    "proposito": {
        "proposito_texto": "Los estudiantes aprenderán a resolver problemas cotidianos utilizando sistemas de ecuaciones lineales con dos incógnitas mediante el método algebraico.",
        "conocimientos": "Sistemas de ecuaciones lineales, método de reducción, sustitución e igualación.",
        "competencia": "Resuelve problemas de regularidad, equivalencia y cambio",
        "estandar": "Resuelve problemas referidos a analizar cambios continuos o periódicos, o regularidades entre magnitudes...",
        "desempeno": "Combina y adapta métodos analíticos para resolver sistemas de ecuaciones lineales.",
        "capacidades": [
            "Traduce datos y condiciones a expresiones algebraicas y gráficas.",
            "Comunica su comprensión sobre las relaciones algebraicas.",
            "Usa estrategias y procedimientos para encontrar equivalencias y reglas generales."
        ],
        "criterios": [
            "Identifica las incógnitas y relaciones en un enunciado problemático.",
            "Aplica el método de reducción para hallar el conjunto solución.",
            "Verifica la validez de los resultados en el contexto del problema."
        ],
        "producto_evidencia": "Ficha de trabajo resuelta con justificación matemática de los métodos empleados.",
        "instrumento": "Lista de Cotejo"
    },
    "competencias_transversales": [
        {
            "titulo": "Gestiona su aprendizaje de manera autónoma",
            "desempenos": [
                "Determina metas de aprendizaje viables asociadas a sus conocimientos y habilidades.",
                "Organiza un conjunto de estrategias y procedimientos en función del tiempo y recursos."
            ]
        },
        {
            "titulo": "Se desenvuelve en entornos virtuales generados por las TIC",
            "desempenos": [
                "Navega en plataformas digitales para contrastar información matemática."
            ]
        }
    ],
    "enfoques_transversales": [
        {
            "nombre": "Enfoque Búsqueda de la Excelencia",
            "valor": "Superación personal",
            "actitudes": "Disposición para adaptarse a los cambios y alcanzar metas complejas con optimismo."
        },
        {
            "nombre": "Enfoque Ambiental",
            "valor": "Justicia y solidaridad",
            "actitudes": "Disposición a evaluar el impacto de las decisiones en el entorno escolar."
        }
    ],
    "recursos": {
        "enlaces": "https://www.perueduca.pe/#/materiales-secundaria-matematica",
        "materiales": "Ficha de actividades N° 04, pizarra, plumones, regla milimetrada.",
        "refuerzo": "Ficha de Refuerzo N° 04: Sistemas lineales"
    },
    "momentos": {
        "inicio": {
            "tiempo_total": "15",
            "actividades": [
                "El docente saluda cordialmente a los estudiantes y acuerda las normas de convivencia.",
                "Se presenta una situación problemática sobre la compra de insumos escolares para el laboratorio.",
                "Se recogen saberes previos mediante preguntas: ¿Qué es una incógnita? ¿Cómo representamos dos cantidades desconocidas?"
            ]
        },
        "desarrollo": {
            "tiempo_total": "65",
            "procesos": [
                {
                    "clave": "familiarizacion",
                    "titulo": "Familiarización con el problema",
                    "contenido": [
                        "Los estudiantes leen en parejas la situación planteada y subrayan los datos numéricos y condiciones.",
                        "El docente realiza preguntas de comprensión para asegurar que todos identifiquen qué se pide calcular."
                    ]
                },
                {
                    "clave": "busqueda_estrategias",
                    "titulo": "Búsqueda y ejecución de estrategias",
                    "contenido": [
                        "Los equipos proponen representar el problema mediante un sistema de dos ecuaciones con dos incógnitas.",
                        "Aplican el método de reducción multiplicando las filas para eliminar una de las variables."
                    ]
                },
                {
                    "clave": "socializacion",
                    "titulo": "Socialización de representaciones",
                    "contenido": [
                        "Dos equipos explican en la pizarra sus procedimientos paso a paso y comparan sus resultados.",
                        "Se valida la solución sustituyendo los valores en las ecuaciones originales."
                    ]
                },
                {
                    "clave": "formalizacion",
                    "titulo": "Formalización y reflexión",
                    "contenido": [
                        "El docente sintetiza los pasos clave del método de reducción y las propiedades algebraicas aplicadas."
                    ]
                }
            ]
        },
        "cierre": {
            "tiempo_total": "10",
            "metacognicion": [
                "¿Qué método algebraico resultó más directo para resolver el sistema?",
                "¿Qué dificultades tuvimos al formular las ecuaciones y cómo las superamos?"
            ],
            "evaluacion": [
                "Los estudiantes autoevalúan su aprendizaje usando los criterios de la lista de cotejo."
            ],
            "extension": [
                "Resolver los problemas 3 y 4 de la página 45 del cuaderno de trabajo de Matemática."
            ]
        }
    },
    "alumnos": [
        "ALVAREZ ROJAS, JUAN CARLOS",
        "BARRERA CASTILLO, MARÍA ELENA",
        "CASTILLO HUAMÁN, ANDRÉS",
        "DELGADO MENDOZA, SOFÍA",
        "ESPINOZA PAREDES, KEVIN"
    ]
}

normalized = normalize_sesion_data(raw_payload)
req = SesionAprendizajeRequest(**normalized)
docx_bytes = build_docx_from_json(req)

doc = docx.Document(docx_bytes)
print(f"Total sections: {len(doc.sections)}")
sec = doc.sections[0]
print(f"Page size: {sec.page_width.inches} x {sec.page_height.inches}")
print(f"Margins (in): Top={sec.top_margin.inches:.2f}, Bottom={sec.bottom_margin.inches:.2f}, Left={sec.left_margin.inches:.2f}, Right={sec.right_margin.inches:.2f}")

print(f"\nTotal tables in document: {len(doc.tables)}")
for i, tbl in enumerate(doc.tables):
    tblPr = tbl._tbl.tblPr
    tblInd = tblPr.find(qn('w:tblInd'))
    tblW = tblPr.find(qn('w:tblW'))
    tblGrid = tbl._tbl.find(qn('w:tblGrid'))
    col_widths = []
    if tblGrid is not None:
        for gc in tblGrid.findall(qn('w:gridCol')):
            col_widths.append(int(gc.attrib.get(qn('w:w'), 0)))
    
    ind_val = tblInd.attrib.get(qn('w:w'), 'None') if tblInd is not None else 'None'
    w_val = tblW.attrib.get(qn('w:w'), 'None') if tblW is not None else 'None'
    print(f"Table {i}: rows={len(tbl.rows)}, cols={len(tbl.columns)}, tblInd={ind_val}, tblW={w_val}, sum(grid)={sum(col_widths)}, grid={col_widths[:6]}")

# Verification assertions
assert sec.top_margin.inches == 0.1, f"Top margin should be 0.1 in, got {sec.top_margin.inches}"
assert sec.bottom_margin.inches == 0.1, f"Bottom margin should be 0.1 in, got {sec.bottom_margin.inches}"
assert sec.left_margin.inches == 0.75, f"Left margin should be 0.75 in, got {sec.left_margin.inches}"
assert sec.right_margin.inches == 0.75, f"Right margin should be 0.75 in, got {sec.right_margin.inches}"

# Table 0 (Datos Informativos)
tbl0 = doc.tables[0]
assert len(tbl0.columns) == 13, f"Table 0 should have 13 columns, got {len(tbl0.columns)}"
assert len(tbl0.rows) == 3, f"Table 0 should have 3 rows, got {len(tbl0.rows)}"

# Table 1 (Título, Propósito, Conocimientos)
tbl1 = doc.tables[1]
assert len(tbl1.columns) == 1, f"Table 1 should have 1 column, got {len(tbl1.columns)}"
assert len(tbl1.rows) == 6, f"Table 1 should have 6 rows, got {len(tbl1.rows)}"

# Table 2 (Propósitos de Aprendizaje)
tbl2 = doc.tables[2]
assert len(tbl2.columns) == 5, f"Table 2 should have 5 columns, got {len(tbl2.columns)}"
assert len(tbl2.rows) == 6, f"Table 2 should have 6 rows, got {len(tbl2.rows)}"

# Table 3 (Competencias Transversales)
tbl3 = doc.tables[3]
assert len(tbl3.columns) == 2, f"Table 3 should have 2 columns, got {len(tbl3.columns)}"

# Table 4 (Enfoques Transversales)
tbl4 = doc.tables[4]
assert len(tbl4.columns) == 3, f"Table 4 should have 3 columns, got {len(tbl4.columns)}"

# Table 5 (Recursos)
tbl5 = doc.tables[5]
assert len(tbl5.columns) == 2, f"Table 5 should have 2 columns, got {len(tbl5.columns)}"
assert len(tbl5.rows) == 4, f"Table 5 should have 4 rows, got {len(tbl5.rows)}"

# Table 6 (Momentos)
tbl6 = doc.tables[6]
assert len(tbl6.columns) == 4, f"Table 6 should have 4 columns, got {len(tbl6.columns)}"

print("\n>>> ALL DOCX FIDELITY CHECKS PASSED WITH 100% SUCCESS! <<<")
