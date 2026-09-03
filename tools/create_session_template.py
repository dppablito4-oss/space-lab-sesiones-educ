"""Create the sanitized DOCX base used by the local session exporter.

The source document is a visual reference supplied by the project.  This
script deliberately removes every body element and body image relationship so
sample names, activities and student rosters can never leak into exports or
the packaged executable.  Header, page geometry, styles, theme and numbering
remain source-derived.
"""
from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
DRAWING_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def _extract_header_banner(source: Path) -> bytes:
    """Extract the wide official banner used by the reference header."""
    with ZipFile(source) as package:
        header = ET.fromstring(package.read("word/header1.xml"))
        rels = ET.fromstring(package.read("word/_rels/header1.xml.rels"))
        targets = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels.findall(f"{{{REL_NS}}}Relationship")
            if rel.attrib.get("Type", "").endswith("/image")
        }
        candidates = []
        for drawing in header.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            extent = drawing.find(".//wp:extent", DRAWING_NS)
            blip = drawing.find(".//a:blip", DRAWING_NS)
            if extent is None or blip is None:
                continue
            rel_id = blip.attrib.get(f"{{{DRAWING_NS['r']}}}embed")
            width = int(extent.attrib.get("cx", "0"))
            height = max(1, int(extent.attrib.get("cy", "1")))
            if rel_id in targets:
                candidates.append((width / height, width, targets[rel_id]))
        if not candidates:
            raise RuntimeError("No se encontró el banner institucional del encabezado.")
        target = max(candidates)[2].replace("\\", "/")
        return package.read(f"word/{target}")


def build_sanitized_template(source: Path, output: Path) -> None:
    banner = _extract_header_banner(source)
    document = Document(source)
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    # Body images are samples. Header images live in the header part and are
    # intentionally retained.
    for rel_id, relationship in list(document.part.rels.items()):
        if relationship.reltype == RT.IMAGE:
            del document.part.rels[rel_id]

    # Replace the source's floating header group with one stable inline banner.
    # It preserves the same official artwork but avoids Word/LibreOffice anchor
    # differences and removes unused header images.
    header = document.sections[0].header
    for child in list(header._element):
        header._element.remove(child)
    for rel_id, relationship in list(header.part.rels.items()):
        if relationship.reltype == RT.IMAGE:
            del header.part.rels[rel_id]
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(
        BytesIO(banner), width=Inches(6.97), height=Inches(0.60)
    )

    properties = document.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.title = "Plantilla oficial de sesión de aprendizaje v1"
    properties.subject = "Base sanitizada para exportaciones DOCX"
    properties.comments = ""
    properties.keywords = ""

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    if args.source.resolve() == args.output.resolve():
        raise SystemExit("La salida debe ser distinta de la plantilla original.")
    build_sanitized_template(args.source, args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
