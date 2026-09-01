"""
DOCX Builder for SessionDocument v1 (Native Pydantic v1 model).

Genera un documento .docx nativo oficial MINEDU con:
  - Formato XML de 13 columnas para datos informativos
  - Tablas con bordes negros 0.5pt (sz=4), colores de encabezado (Peach, Blue, Yellow)
  - Procesos pedagógicos ordenados con contenido HTML/Rich
  - Ficha de trabajo y Lista de cotejo dinámicas
  - Soporte para Educación Inicial (Juego Libre en los Sectores)
"""
from __future__ import annotations
import io
import re
import sys
from pathlib import Path
from typing import List, Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from models.session_document import (
    SessionDocumentV1, MetadataV1, PropositoV1,
    CompetenciaTransversalV1, EnfoqueTransversalV1,
    RecursosV1, MomentosV1, MomentoV1, SessionProcess,
    EvaluacionV1, FichaTrabajoV1, JuegoLibreSectoresV1, ListaCotejoV1
)
from docx_builder import (
    set_cell_background,
    set_cell_margins,
    add_table_borders,
    add_table_borders_black,
    set_table_col_widths_and_indent,
    set_cell_text_white_bold,
    format_latex_to_unicode,
    append_html_to_cell_or_paragraph
)


def _hex(value: str, fallback: str) -> str:
    clean = str(value or '').lstrip('#').upper()
    return clean if re.fullmatch(r'[0-9A-F]{6}', clean) else fallback


def _blend_with_white(hex_color: str, white_ratio: float) -> str:
    ratio = max(0.0, min(1.0, white_ratio))
    rgb = [int(hex_color[i:i + 2], 16) for i in (0, 2, 4)]
    return ''.join(f'{round(channel * (1 - ratio) + 255 * ratio):02X}' for channel in rgb)


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor(*(int(hex_color[i:i + 2], 16) for i in (0, 2, 4)))


def _apply_presentation(doc: Document, doc_v1: SessionDocumentV1) -> None:
    """Apply final shared tokens without flattening the document structure."""
    p = doc_v1.presentation
    primary = _hex(p.primaryColor, '000000')
    scale = float(p.fontSizePt) / 10.0
    margins = {
        'compact': (40, 70), 'standard': (70, 100),
        'comfortable': (100, 140), 'spacious': (130, 180),
    }
    vertical, horizontal = margins[p.cellPadding]

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = float(p.lineHeight)
        for run in paragraph.runs:
            run.font.name = p.fontFamily
            run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), p.fontFamily)
            if run.font.size:
                run.font.size = Pt(max(7, run.font.size.pt * scale))

    for table in doc.tables:
        for edge in table._tbl.xpath('.//w:tblBorders/*'):
            edge.set(qn('w:color'), primary)
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=vertical, bottom=vertical, left=horizontal, right=horizontal)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = float(p.lineHeight)
                    for run in paragraph.runs:
                        run.font.name = p.fontFamily
                        run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), p.fontFamily)
                        if run.font.size:
                            run.font.size = Pt(max(7, run.font.size.pt * scale))

def _preprocess_v1_latex(doc_v1: SessionDocumentV1):
    """Preprocesa todas las cadenas del documento v1 convirtiendo LaTeX a Unicode."""
    if doc_v1.metadata.titulo:
        doc_v1.metadata.titulo = format_latex_to_unicode(doc_v1.metadata.titulo)
    if doc_v1.proposito.texto:
        doc_v1.proposito.texto = format_latex_to_unicode(doc_v1.proposito.texto)
    if doc_v1.proposito.conocimientos:
        doc_v1.proposito.conocimientos = format_latex_to_unicode(doc_v1.proposito.conocimientos)
    if doc_v1.proposito.estandar:
        doc_v1.proposito.estandar = format_latex_to_unicode(doc_v1.proposito.estandar)
    if doc_v1.proposito.evidencia:
        doc_v1.proposito.evidencia = format_latex_to_unicode(doc_v1.proposito.evidencia)

    doc_v1.proposito.capacidades = [format_latex_to_unicode(c) for c in doc_v1.proposito.capacidades]
    doc_v1.proposito.criterios = [format_latex_to_unicode(c) for c in doc_v1.proposito.criterios]

    for m in [doc_v1.momentos.inicio, doc_v1.momentos.desarrollo, doc_v1.momentos.cierre]:
        for proc in m.procesos:
            if proc.titulo:
                proc.titulo = format_latex_to_unicode(proc.titulo)
            if proc.contenido and proc.contenido.value:
                proc.contenido.value = format_latex_to_unicode(proc.contenido.value)

    if doc_v1.fichaTrabajo:
        if doc_v1.fichaTrabajo.titulo:
            doc_v1.fichaTrabajo.titulo = format_latex_to_unicode(doc_v1.fichaTrabajo.titulo)
        if doc_v1.fichaTrabajo.indicaciones:
            doc_v1.fichaTrabajo.indicaciones = format_latex_to_unicode(doc_v1.fichaTrabajo.indicaciones)
        if doc_v1.fichaTrabajo.actividades:
            doc_v1.fichaTrabajo.actividades = format_latex_to_unicode(doc_v1.fichaTrabajo.actividades)


def build_docx_from_v1(doc_v1: SessionDocumentV1) -> io.BytesIO:
    """
    Construye un documento .docx nativo oficial a partir de un SessionDocumentV1.
    """
    _preprocess_v1_latex(doc_v1)

    doc = Document()

    # Márgenes exactos A4
    for s in doc.sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(0.1)
        s.bottom_margin = Inches(0.1)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    presentation = doc_v1.presentation
    PRIMARY = _hex(presentation.primaryColor, '000000')
    ACCENT = _hex(presentation.accentColor, 'C0392B')
    HEADER_BG = _hex(presentation.headerBackground, 'BDD6EE')
    PRIMARY_RGB = _rgb(PRIMARY)
    ACCENT_RGB = _rgb(ACCENT)
    style_normal.font.name = presentation.fontFamily
    style_normal.font.size = Pt(presentation.fontSizePt)
    style_normal.font.color.rgb = PRIMARY_RGB

    # Paleta de colores institucional
    PEACH = _blend_with_white(ACCENT, 0.78)
    BLUE_HDR = HEADER_BG
    YELLOW_HDR = _blend_with_white(PRIMARY, 0.82)
    GRAY_VAL = _blend_with_white(PRIMARY, 0.94)
    PEACH_MOM = _blend_with_white(ACCENT, 0.88)
    GRAY_MOM = _blend_with_white(PRIMARY, 0.90)
    YELLOW_VAL = _blend_with_white(ACCENT, 0.92)
    BULLET_COLORS = [ACCENT, PRIMARY, '277A4B', '6D3A8C', '0F766E']

    def _label(cell, text):
        set_cell_background(cell, PEACH)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)

    def _val(cell, text):
        set_cell_background(cell, GRAY_VAL)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        p.add_run(text or "").font.size = Pt(9)

    def _hdr(cell, text, bg=BLUE_HDR, sz=9):
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(sz)

    def _bullet_cell(cell, items):
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        for i, item in enumerate(items):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            bc = BULLET_COLORS[i % len(BULLET_COLORS)]
            rb = p.add_run(u"\u25cf ")
            rb.font.size = Pt(9)
            rb.font.color.rgb = _rgb(bc)
            append_html_to_cell_or_paragraph(p, item, default_font_size=9)

    def make_vertical_text(text: str) -> str:
        return "\n" + "\n\n".join(list(text)) + "\n"

    def _write_momento_cell(cell, nombre, sub_bullets, tiempo):
        set_cell_background(cell, GRAY_MOM)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p_main = cell.paragraphs[0]
        rm = p_main.add_run(nombre)
        rm.bold = True
        rm.font.size = Pt(10)
        if sub_bullets:
            p_sub = cell.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(4)
            rs = p_sub.add_run(sub_bullets)
            rs.font.size = Pt(8)
            rs.font.color.rgb = RGBColor(71, 85, 105)
        if tiempo:
            p_t = cell.add_paragraph()
            p_t.paragraph_format.space_before = Pt(6)
            rt = p_t.add_run(u"\u23f1 TIEMPO: " + str(tiempo) + " min")
            rt.bold = True
            rt.font.size = Pt(8.5)
            rt.font.color.rgb = RGBColor(192, 57, 43)

    def _write_vertical_cell(cell, txt):
        set_cell_background(cell, PEACH_MOM)
        set_cell_margins(cell, top=120, bottom=120, left=40, right=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(make_vertical_text(txt))
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = ACCENT_RGB

    # ── Header nativo ──
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False

    logo_stream = None
    try:
        logo_path = Path(__file__).resolve().parent.parent / "assets" / "logo.png"
        if logo_path.exists():
            with open(logo_path, "rb") as f:
                logo_stream = io.BytesIO(f.read())
    except Exception:
        pass

    p_logo = header.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(4)

    if logo_stream:
        try:
            p_logo.add_run().add_picture(logo_stream, height=Inches(0.55))
        except Exception:
            lbl = p_logo.add_run("🚀 Space Lab")
            lbl.bold = True
            lbl.font.size = Pt(11)
            lbl.font.color.rgb = RGBColor(56, 189, 248)
    else:
        lbl = p_logo.add_run("🚀 Space Lab")
        lbl.bold = True
        lbl.font.size = Pt(11)
        lbl.font.color.rgb = RGBColor(56, 189, 248)

    p_div = header.add_paragraph()
    p_div.paragraph_format.space_before = Pt(3)
    p_div.paragraph_format.space_after = Pt(6)
    pBrd = OxmlElement('w:pBrd')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '12')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '334155')
    pBrd.append(bottom_border)
    p_div._p.get_or_add_pPr().append(pBrd)

    meta = doc_v1.metadata
    prop = doc_v1.proposito

    # The reference template starts with the official session identifier before
    # the informative-data table. Keep this visible in every DOCX export.
    p_session = doc.add_paragraph()
    p_session.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_session.paragraph_format.space_before = Pt(0)
    p_session.paragraph_format.space_after = Pt(5)
    r_session = p_session.add_run(
        f"SESIÓN DE APRENDIZAJE N° {meta.numeroSesion or '01'}"
    )
    r_session.bold = True
    r_session.font.size = Pt(12)
    r_session.font.color.rgb = RGBColor(0, 0, 0)

    # ── TABLA 0: DATOS INFORMATIVOS (13 columnas, 3 filas) ──
    _DI_TWIPS = [1221, 622, 391, 601, 1134, 850, 851, 568, 142, 1275, 426, 1134, 1274]
    di = doc.add_table(rows=3, cols=13)
    di.autofit = False
    add_table_borders_black(di, sz='4')

    # Fila 0
    c_ie_lbl = di.cell(0, 0).merge(di.cell(0, 2))
    _label(c_ie_lbl, "Institución Educativa")
    c_ie_val = di.cell(0, 3).merge(di.cell(0, 6))
    _val(c_ie_val, meta.institucion or "I.E.")

    c_nv_lbl = di.cell(0, 7).merge(di.cell(0, 8))
    _label(c_nv_lbl, "Nivel")
    c_nv_val = di.cell(0, 9).merge(di.cell(0, 12))
    _val(c_nv_val, meta.nivel or "SECUNDARIA")

    # Fila 1
    _label(di.cell(1, 0), "Docente")
    c_doc_val = di.cell(1, 1).merge(di.cell(1, 6))
    _val(c_doc_val, meta.docente or "")

    c_ar_lbl = di.cell(1, 7).merge(di.cell(1, 8))
    _label(c_ar_lbl, "Área")
    c_ar_val = di.cell(1, 9).merge(di.cell(1, 10))
    _val(c_ar_val, meta.area or "")

    _label(di.cell(1, 11), "Unidad/ Proyecto")
    _val(di.cell(1, 12), meta.unidad or "")

    # Fila 2
    _label(di.cell(2, 0), "Grado")
    _val(di.cell(2, 1), meta.grado or "")

    c_sec_lbl = di.cell(2, 2).merge(di.cell(2, 3))
    _label(c_sec_lbl, "Sección")
    _val(di.cell(2, 4), meta.seccion or "")

    _label(di.cell(2, 5), "Fecha")
    c_fec_val = di.cell(2, 6).merge(di.cell(2, 7))
    _val(c_fec_val, meta.fecha or "")

    c_dur_lbl = di.cell(2, 8).merge(di.cell(2, 9))
    _label(c_dur_lbl, "Duración (minutos)")
    c_dur_val = di.cell(2, 10).merge(di.cell(2, 12))
    _val(c_dur_val, f"{meta.duracionMinutos} min" if meta.duracionMinutos else "90 min")

    set_table_col_widths_and_indent(di, _DI_TWIPS, indent_twip=-289)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ── TABLA 1: TÍTULO, PROPÓSITO Y CONOCIMIENTOS ──
    _T1_TWIPS = [10490]
    pc = doc.add_table(rows=6, cols=1)
    pc.autofit = False
    add_table_borders_black(pc, sz='4')

    _hdr(pc.cell(0, 0), "TÍTULO DE LA SESIÓN", bg=BLUE_HDR, sz=9.5)

    cell_tit = pc.cell(1, 0)
    set_cell_background(cell_tit, 'FFFFFF')
    set_cell_margins(cell_tit, top=80, bottom=80, left=120, right=120)
    p_tit = cell_tit.paragraphs[0]
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_before = Pt(0)
    p_tit.paragraph_format.space_after = Pt(0)
    p_tit.paragraph_format.line_spacing = 1.1
    r_tit = p_tit.add_run(meta.titulo or "Título de la sesión de aprendizaje")
    r_tit.bold = True
    r_tit.font.size = Pt(10.5)
    r_tit.font.color.rgb = ACCENT_RGB

    _hdr(pc.cell(2, 0), "PROPÓSITO DE LA SESIÓN", bg=BLUE_HDR, sz=9.5)

    cell_prop = pc.cell(3, 0)
    set_cell_background(cell_prop, 'FFFFFF')
    set_cell_margins(cell_prop, top=80, bottom=80, left=120, right=120)
    p_prop = cell_prop.paragraphs[0]
    p_prop.paragraph_format.space_before = Pt(0)
    p_prop.paragraph_format.space_after = Pt(0)
    p_prop.paragraph_format.line_spacing = 1.15
    append_html_to_cell_or_paragraph(p_prop, prop.texto or "No especificado", default_font_size=9.5)

    _hdr(pc.cell(4, 0), "CONOCIMIENTOS", bg=YELLOW_HDR, sz=9.5)

    cell_con = pc.cell(5, 0)
    set_cell_background(cell_con, 'FFFFFF')
    set_cell_margins(cell_con, top=80, bottom=80, left=120, right=120)
    p_con = cell_con.paragraphs[0]
    p_con.paragraph_format.space_before = Pt(0)
    p_con.paragraph_format.space_after = Pt(0)
    p_con.paragraph_format.line_spacing = 1.15
    append_html_to_cell_or_paragraph(p_con, prop.conocimientos or "No especificado", default_font_size=9.5)

    set_table_col_widths_and_indent(pc, _T1_TWIPS, indent_twip=-289)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ── TABLA 2: PROPÓSITOS DE APRENDIZAJE ──
    _PA_TWIPS = [1702, 2410, 3543, 1418, 1417]
    pa = doc.add_table(rows=6, cols=5)
    pa.autofit = False
    add_table_borders_black(pa, sz='4')

    c_hdr = pa.cell(0, 0).merge(pa.cell(0, 4))
    _hdr(c_hdr, "PROPÓSITOS DE APRENDIZAJE", bg=BLUE_HDR, sz=10)

    c_comp = pa.cell(1, 0).merge(pa.cell(1, 4))
    set_cell_background(c_comp, 'FFFFFF')
    set_cell_margins(c_comp, top=80, bottom=80, left=140, right=140)
    p_comp = c_comp.paragraphs[0]
    p_comp.paragraph_format.space_before = Pt(0)
    p_comp.paragraph_format.space_after = Pt(0)
    p_comp.paragraph_format.line_spacing = 1.15
    rc1 = p_comp.add_run("Competencia: ")
    rc1.bold = True
    rc1.font.size = Pt(9.5)
    rc2 = p_comp.add_run(prop.competencia or "No especificada")
    rc2.font.size = Pt(9.5)

    c_est = pa.cell(2, 0).merge(pa.cell(2, 4))
    set_cell_background(c_est, 'FFFFFF')
    set_cell_margins(c_est, top=80, bottom=80, left=140, right=140)
    p_est = c_est.paragraphs[0]
    p_est.paragraph_format.space_before = Pt(0)
    p_est.paragraph_format.space_after = Pt(0)
    p_est.paragraph_format.line_spacing = 1.15
    re1 = p_est.add_run("Estándar de aprendizaje: ")
    re1.bold = True
    re1.font.size = Pt(9.5)
    re2 = p_est.add_run(prop.estandar or "No especificado")
    re2.font.size = Pt(9.5)

    headers_pa = ["COMPETENCIAS", "CAPACIDADES", "CRITERIOS DE EVALUACION", "PRODUCTO / EVIDENCIA", "INSTRUMENTOS DE EVALUACIÓN"]
    for i, ht in enumerate(headers_pa):
        _hdr(pa.cell(3, i), ht, bg=BLUE_HDR, sz=8.0)

    c_comp_v = pa.cell(4, 0).merge(pa.cell(5, 0))
    set_cell_background(c_comp_v, 'FFFFFF')
    set_cell_margins(c_comp_v, top=100, bottom=100, left=100, right=100)
    p_cv = c_comp_v.paragraphs[0]
    p_cv.paragraph_format.space_before = Pt(0)
    p_cv.paragraph_format.space_after = Pt(0)
    p_cv.paragraph_format.line_spacing = 1.1
    p_cv.add_run(prop.competencia or "No especificada").font.size = Pt(8.5)

    _bullet_cell(pa.cell(4, 1), prop.capacidades)
    _bullet_cell(pa.cell(4, 2), prop.criterios)

    c_ev = pa.cell(4, 3)
    set_cell_background(c_ev, 'FFFFFF')
    set_cell_margins(c_ev, top=100, bottom=100, left=100, right=100)
    p_ev = c_ev.paragraphs[0]
    p_ev.paragraph_format.space_before = Pt(0)
    p_ev.paragraph_format.space_after = Pt(0)
    p_ev.paragraph_format.line_spacing = 1.1
    p_ev.add_run(prop.evidencia or "No especificado").font.size = Pt(8.5)

    c_ins_v = pa.cell(4, 4).merge(pa.cell(5, 4))
    set_cell_background(c_ins_v, GRAY_VAL)
    set_cell_margins(c_ins_v, top=100, bottom=100, left=100, right=100)
    p_iv = c_ins_v.paragraphs[0]
    p_iv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_iv.paragraph_format.space_before = Pt(0)
    p_iv.paragraph_format.space_after = Pt(0)
    p_iv.paragraph_format.line_spacing = 1.1
    ri_v = p_iv.add_run(prop.instrumento or "Lista de Cotejo")
    ri_v.bold = True
    ri_v.font.size = Pt(8.5)

    set_cell_margins(pa.cell(5, 1), top=40, bottom=40, left=40, right=40)
    set_cell_margins(pa.cell(5, 2), top=40, bottom=40, left=40, right=40)
    set_cell_margins(pa.cell(5, 3), top=40, bottom=40, left=40, right=40)

    set_table_col_widths_and_indent(pa, _PA_TWIPS, indent_twip=-289)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ── TABLA 3: COMPETENCIAS TRANSVERSALES ──
    _CT_TWIPS = [3403, 7087]
    cts = doc_v1.competenciasTransversales
    n_ct = len(cts) if cts else 2
    ct_tbl = doc.add_table(rows=1 + n_ct, cols=2)
    ct_tbl.autofit = False
    add_table_borders_black(ct_tbl, sz='4')

    _hdr(ct_tbl.cell(0, 0), "COMPETENCIAS TRANSVERSALES", bg=YELLOW_HDR, sz=9)
    _hdr(ct_tbl.cell(0, 1), "CRITERIOS DE EVALUACION", bg=YELLOW_HDR, sz=9)

    if cts:
        for ci, ct in enumerate(cts):
            rn = 1 + ci
            set_cell_background(ct_tbl.cell(rn, 0), GRAY_VAL)
            set_cell_margins(ct_tbl.cell(rn, 0), top=80, bottom=80, left=120, right=120)
            p_ct = ct_tbl.cell(rn, 0).paragraphs[0]
            p_ct.paragraph_format.space_before = Pt(0)
            p_ct.paragraph_format.space_after = Pt(0)
            p_ct.paragraph_format.line_spacing = 1.1
            p_ct.add_run(ct.titulo or "Competencia Transversal").font.size = Pt(8.5)
            _bullet_cell(ct_tbl.cell(rn, 1), ct.desempenos if ct.desempenos else ["No especificado"])
    else:
        defaults_ct = [
            ("Gestiona su aprendizaje de manera autonoma", [
                "Define metas de aprendizaje para alcanzar sus objetivos pedagógicos.",
                "Organiza acciones estrategicas para alcanzar sus metas de aprendizaje.",
                "Monitorea y ajusta su desempeno durante el proceso de aprendizaje."
            ]),
            ("Se desenvuelve en los entornos virtuales generados por las TIC", [
                "Personaliza entornos virtuales segun sus necesidades de indagacion.",
                "Gestiona informacion del entorno virtual de manera segura.",
                "Interactua en entornos virtuales y crea objetos virtuales."
            ])
        ]
        for ci, (title, items) in enumerate(defaults_ct):
            rn = 1 + ci
            set_cell_background(ct_tbl.cell(rn, 0), GRAY_VAL)
            set_cell_margins(ct_tbl.cell(rn, 0), top=80, bottom=80, left=120, right=120)
            p_ct = ct_tbl.cell(rn, 0).paragraphs[0]
            p_ct.paragraph_format.space_before = Pt(0)
            p_ct.paragraph_format.space_after = Pt(0)
            p_ct.paragraph_format.line_spacing = 1.1
            p_ct.add_run(title).font.size = Pt(8.5)
            _bullet_cell(ct_tbl.cell(rn, 1), items)

    set_table_col_widths_and_indent(ct_tbl, _CT_TWIPS, indent_twip=-289)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ── TABLA 4: ENFOQUES TRANSVERSALES ──
    _ET_TWIPS = [2269, 1559, 6662]
    enfoques = doc_v1.enfoquesTransversales
    n_enf = len(enfoques) if enfoques else 2
    et_tbl = doc.add_table(rows=1 + n_enf, cols=3)
    et_tbl.autofit = False
    add_table_borders_black(et_tbl, sz='4')

    headers_et = ["Enfoque(s) transversal(es)", "Valores", "Actitudes o acciones observables"]
    for i, ht in enumerate(headers_et):
        _hdr(et_tbl.cell(0, i), ht, bg=YELLOW_HDR, sz=9)

    if enfoques:
        for ci, et in enumerate(enfoques):
            rn = 1 + ci
            set_cell_background(et_tbl.cell(rn, 0), GRAY_VAL)
            set_cell_margins(et_tbl.cell(rn, 0), top=80, bottom=80, left=120, right=120)
            p_e0 = et_tbl.cell(rn, 0).paragraphs[0]
            p_e0.paragraph_format.space_before = Pt(0)
            p_e0.paragraph_format.space_after = Pt(0)
            p_e0.paragraph_format.line_spacing = 1.1
            p_e0.add_run(et.nombre or "Enfoque Transversal").font.size = Pt(8.5)

            set_cell_background(et_tbl.cell(rn, 1), YELLOW_VAL)
            set_cell_margins(et_tbl.cell(rn, 1), top=80, bottom=80, left=120, right=120)
            p_e1 = et_tbl.cell(rn, 1).paragraphs[0]
            p_e1.paragraph_format.space_before = Pt(0)
            p_e1.paragraph_format.space_after = Pt(0)
            p_e1.paragraph_format.line_spacing = 1.1
            p_e1.add_run(et.valor or "No especificado").font.size = Pt(8.5)

            set_cell_background(et_tbl.cell(rn, 2), YELLOW_VAL)
            set_cell_margins(et_tbl.cell(rn, 2), top=80, bottom=80, left=120, right=120)
            p_e2 = et_tbl.cell(rn, 2).paragraphs[0]
            p_e2.paragraph_format.space_before = Pt(0)
            p_e2.paragraph_format.space_after = Pt(0)
            p_e2.paragraph_format.line_spacing = 1.1
            p_e2.add_run(et.actitudes or "No especificadas").font.size = Pt(8.5)
    else:
        defaults_et = [
            ("Enfoque Ambiental", "Justicia y solidaridad", "Reduce el uso de materiales desechables, reutilizando cuadernos, hojas y envases cuando sea posible durante las actividades del aula."),
            ("Enfoque Busqueda de la Excelencia", "Equidad y Justicia", "Dialoga con tus compañeros para resolver desacuerdos y escucha con atencion.")
        ]
        for ci, (name, val, act) in enumerate(defaults_et):
            rn = 1 + ci
            set_cell_background(et_tbl.cell(rn, 0), GRAY_VAL)
            set_cell_margins(et_tbl.cell(rn, 0), top=80, bottom=80, left=120, right=120)
            p_e0 = et_tbl.cell(rn, 0).paragraphs[0]
            p_e0.paragraph_format.space_before = Pt(0)
            p_e0.paragraph_format.space_after = Pt(0)
            p_e0.paragraph_format.line_spacing = 1.1
            p_e0.add_run(name).font.size = Pt(8.5)

            set_cell_background(et_tbl.cell(rn, 1), YELLOW_VAL)
            set_cell_margins(et_tbl.cell(rn, 1), top=80, bottom=80, left=120, right=120)
            p_e1 = et_tbl.cell(rn, 1).paragraphs[0]
            p_e1.paragraph_format.space_before = Pt(0)
            p_e1.paragraph_format.space_after = Pt(0)
            p_e1.paragraph_format.line_spacing = 1.1
            p_e1.add_run(val).font.size = Pt(8.5)

            set_cell_background(et_tbl.cell(rn, 2), YELLOW_VAL)
            set_cell_margins(et_tbl.cell(rn, 2), top=80, bottom=80, left=120, right=120)
            p_e2 = et_tbl.cell(rn, 2).paragraphs[0]
            p_e2.paragraph_format.space_before = Pt(0)
            p_e2.paragraph_format.space_after = Pt(0)
            p_e2.paragraph_format.line_spacing = 1.1
            p_e2.add_run(act).font.size = Pt(8.5)

    set_table_col_widths_and_indent(et_tbl, _ET_TWIPS, indent_twip=-289)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ── TABLA 5: RECURSOS Y MATERIALES ──
    _REC_TWIPS = [3403, 7087]
    rec = doc.add_table(rows=4, cols=2)
    rec.autofit = False
    add_table_borders_black(rec, sz='4')

    _hdr(rec.cell(0, 0), "Páginas de: Texto de, otros textos de consulta/ Enlace web, etc.", bg=YELLOW_HDR, sz=8.5)
    set_cell_background(rec.cell(0, 1), GRAY_VAL)
    set_cell_margins(rec.cell(0, 1), top=80, bottom=80, left=120, right=120)
    p_r0 = rec.cell(0, 1).paragraphs[0]
    p_r0.paragraph_format.space_before = Pt(0)
    p_r0.paragraph_format.space_after = Pt(0)
    p_r0.paragraph_format.line_spacing = 1.1
    p_r0.add_run(doc_v1.recursos.enlaces or "https://www.perueduca.pe/#/home/materiales-educativos").font.size = Pt(8.5)

    _hdr(rec.cell(1, 0), "", bg=YELLOW_HDR, sz=8.5)
    set_cell_background(rec.cell(1, 1), GRAY_VAL)
    set_cell_margins(rec.cell(1, 1), top=40, bottom=40, left=120, right=120)

    _hdr(rec.cell(2, 0), "Materiales y recursos", bg=PEACH_MOM, sz=8.5)
    set_cell_background(rec.cell(2, 1), GRAY_VAL)
    set_cell_margins(rec.cell(2, 1), top=80, bottom=80, left=120, right=120)
    p_r2 = rec.cell(2, 1).paragraphs[0]
    p_r2.paragraph_format.space_before = Pt(0)
    p_r2.paragraph_format.space_after = Pt(0)
    p_r2.paragraph_format.line_spacing = 1.1
    p_r2.add_run(doc_v1.recursos.materiales or "Ficha de actividades N° 01-02").font.size = Pt(8.5)

    _hdr(rec.cell(3, 0), "Actividades de Refuerzo Escolar (N° ficha y Título)", bg=YELLOW_HDR, sz=8.5)
    set_cell_background(rec.cell(3, 1), GRAY_VAL)
    set_cell_margins(rec.cell(3, 1), top=80, bottom=80, left=120, right=120)
    p_r3 = rec.cell(3, 1).paragraphs[0]
    p_r3.paragraph_format.space_before = Pt(0)
    p_r3.paragraph_format.space_after = Pt(0)
    p_r3.paragraph_format.line_spacing = 1.1
    p_r3.add_run(doc_v1.recursos.refuerzo or "No especificado").font.size = Pt(8.5)

    set_table_col_widths_and_indent(rec, _REC_TWIPS, indent_twip=-289)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # ── TABLA 6: SECUENCIAL DE MOMENTOS DE LA SESIÓN ──
    _MOM_TWIPS = [1702, 283, 8080, 425]
    des_procs = doc_v1.momentos.desarrollo.procesos
    n_proc = len(des_procs) if des_procs else 1
    n_rows = 1 + 1 + n_proc + 1

    mt = doc.add_table(rows=n_rows, cols=4)
    mt.autofit = False
    add_table_borders_black(mt, sz='4')

    _hdr(mt.cell(0, 0), "MOMENTOS DE LA SESIÓN", bg=YELLOW_HDR, sz=9.5)
    cell_est_hdr = mt.cell(0, 1).merge(mt.cell(0, 3))
    _hdr(cell_est_hdr, "ESTRATEGIAS / ACTIVIDADES", bg=YELLOW_HDR, sz=9.5)

    # Fila 1: INICIO
    _write_momento_cell(mt.cell(1, 0), "INICIO:",
        "Saberes Previos\nProblematización\nPropósito y organización",
        str(doc_v1.momentos.inicio.tiempoMinutos))
    _write_vertical_cell(mt.cell(1, 1), "MOTIVACION")
    _write_vertical_cell(mt.cell(1, 3), "EVALUACION")

    cell_ini = mt.cell(1, 2)
    set_cell_background(cell_ini, 'FFFFFF')
    set_cell_margins(cell_ini, top=100, bottom=100, left=140, right=140)
    cell_ini.text = ""
    for proc in doc_v1.momentos.inicio.procesos:
        if proc.titulo and proc.id not in ('actividad', 'proceso_1'):
            ppt = cell_ini.add_paragraph()
            ppt.paragraph_format.space_before = Pt(2)
            ppt.paragraph_format.space_after = Pt(2)
            rpt = ppt.add_run(proc.titulo.upper() + ":")
            rpt.bold = True
            rpt.font.size = Pt(9.5)
            rpt.font.color.rgb = PRIMARY_RGB
        append_html_to_cell_or_paragraph(cell_ini, proc.contenido.value, default_font_size=9.5)

    # Filas 2 a 2+n_proc-1: DESARROLLO
    _write_momento_cell(mt.cell(2, 0), "DESARROLLO:",
        "Gestión y Acompañamiento del Desarrollo de las Competencias\n(Procesos didácticos del Área)",
        str(doc_v1.momentos.desarrollo.tiempoMinutos))
    _write_vertical_cell(mt.cell(2, 1), "MOTIVACION")
    _write_vertical_cell(mt.cell(2, 3), "EVALUACION")

    for idx in range(n_proc):
        rn = 2 + idx
        cell_des = mt.cell(rn, 2)
        set_cell_background(cell_des, 'FFFFFF')
        set_cell_margins(cell_des, top=100, bottom=100, left=140, right=140)
        cell_des.text = ""
        if des_procs:
            proc = des_procs[idx]
            ppt = cell_des.add_paragraph()
            ppt.paragraph_format.space_before = Pt(2)
            ppt.paragraph_format.space_after = Pt(4)
            rpt = ppt.add_run(proc.titulo.upper())
            rpt.bold = True
            rpt.font.size = Pt(9.5)
            rpt.font.color.rgb = ACCENT_RGB
            append_html_to_cell_or_paragraph(cell_des, proc.contenido.value, default_font_size=9.5)
        else:
            append_html_to_cell_or_paragraph(cell_des, "Gestión y Acompañamiento del Desarrollo de Competencias...", default_font_size=9.5)

        if idx > 0:
            _write_vertical_cell(mt.cell(rn, 1), "MOTIVACION")
            _write_vertical_cell(mt.cell(rn, 3), "EVALUACION")

    if n_proc > 1:
        mt.cell(2, 0).merge(mt.cell(2 + n_proc - 1, 0))
        mt.cell(2, 1).merge(mt.cell(2 + n_proc - 1, 1))
        mt.cell(2, 3).merge(mt.cell(2 + n_proc - 1, 3))

    # Fila CIERRE
    rc = n_rows - 1
    _write_momento_cell(mt.cell(rc, 0), "CIERRE:",
        "Evaluación (Reflexión sobre lo aprendido)\nAcciones de reforzamiento o indagación",
        str(doc_v1.momentos.cierre.tiempoMinutos))
    _write_vertical_cell(mt.cell(rc, 1), "MOTIVACION")
    _write_vertical_cell(mt.cell(rc, 3), "EVALUACION")

    cell_cie = mt.cell(rc, 2)
    set_cell_background(cell_cie, 'FFFFFF')
    set_cell_margins(cell_cie, top=100, bottom=100, left=140, right=140)
    cell_cie.text = ""

    if doc_v1.momentos.cierre.procesos:
        for proc in doc_v1.momentos.cierre.procesos:
            plbl = cell_cie.add_paragraph()
            plbl.paragraph_format.space_before = Pt(4)
            rlbl = plbl.add_run(proc.titulo + ":")
            rlbl.bold = True
            rlbl.font.size = Pt(9.5)
            append_html_to_cell_or_paragraph(cell_cie, proc.contenido.value, default_font_size=9.5)
    else:
        default_sections = [
            ("Metacognición:", [
                "¿Qué aprendimos hoy? ¿Cómo lo aprendimos? ¿Para qué nos sirve?",
                "¿Qué fue lo más difícil? ¿Cómo lo superamos?"
            ]),
            ("Evaluación formativa:", [
                "Revisión de los criterios de evaluación con los estudiantes.",
                "Retroalimentación sobre el desempeño de la sesión."
            ]),
            ("Extensión para casa:", [
                "Actividad de refuerzo o aplicación a nuevas situaciones.",
                "Resolución de ejercicios complementarios."
            ]),
        ]
        for lbl, items in default_sections:
            plbl = cell_cie.add_paragraph()
            plbl.paragraph_format.space_before = Pt(4)
            rlbl = plbl.add_run(lbl)
            rlbl.bold = True
            rlbl.font.size = Pt(9.5)
            for item in items:
                p = cell_cie.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                rb = p.add_run(u"\u25cf ")
                rb.font.size = Pt(8.5)
                rb.font.color.rgb = PRIMARY_RGB
                append_html_to_cell_or_paragraph(p, item, default_font_size=9.5)

    set_table_col_widths_and_indent(mt, _MOM_TWIPS, indent_twip=-289)

    # ── JUEGO LIBRE EN LOS SECTORES (INICIAL) ──
    jls = doc_v1.juegoLibreSectores
    if jls:
        doc.add_paragraph().paragraph_format.space_before = Pt(8)
        jls_header = doc.add_table(rows=1, cols=1)
        jls_header.autofit = False
        jls_header.rows[0].cells[0].width = Inches(6.77)
        add_table_borders_black(jls_header)
        set_cell_text_white_bold(jls_header.cell(0, 0), "JUEGO LIBRE EN LOS SECTORES", font_size_pt=10)
        set_cell_background(jls_header.cell(0, 0), '27AE60')
        set_cell_margins(jls_header.cell(0, 0), top=100, bottom=100, left=180, right=180)

        jls_steps = [
            ("1. PLANIFICACIÓN", jls.planificacion or "Los niños y niñas eligen libremente el sector donde desean jugar."),
            ("2. ORGANIZACIÓN", jls.organizacion or "Se agrupan según el sector elegido y distribuyen roles."),
            ("3. EJECUCIÓN", jls.ejecucion or "Los niños juegan libremente mientras la docente acompaña y media."),
            ("4. ORDEN", jls.orden or "A la señal, los niños guardan los materiales con una canción motivadora."),
            ("5. SOCIALIZACIÓN", jls.socializacion or "Los niños cuentan lo que hicieron en su sector y qué aprendieron."),
            ("6. REPRESENTACIÓN", jls.representacion or "Los niños dibujan, modelan o dramatizan lo vivido en el juego.")
        ]

        jls_tbl = doc.add_table(rows=len(jls_steps), cols=2)
        jls_tbl.autofit = False
        add_table_borders_black(jls_tbl)

        for ri, (label, content) in enumerate(jls_steps):
            set_cell_background(jls_tbl.cell(ri, 0), 'D5F5E3')
            set_cell_margins(jls_tbl.cell(ri, 0), top=80, bottom=80, left=120, right=120)
            p_lbl = jls_tbl.cell(ri, 0).paragraphs[0]
            r_lbl = p_lbl.add_run(label)
            r_lbl.bold = True
            r_lbl.font.size = Pt(8.5)
            r_lbl.font.color.rgb = RGBColor(30, 41, 59)

            set_cell_background(jls_tbl.cell(ri, 1), 'F2F2F2')
            set_cell_margins(jls_tbl.cell(ri, 1), top=80, bottom=80, left=120, right=120)
            append_html_to_cell_or_paragraph(jls_tbl.cell(ri, 1), content, default_font_size=8.5)

        anchos_jls = [Inches(1.8), Inches(4.97)]
        for row in jls_tbl.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = anchos_jls[ci]

    # ── FIRMAS ──
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    ft = doc.add_table(rows=1, cols=2)
    ft.autofit = True
    for cell in ft.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcB.append(b)
        tcPr.append(tcB)

    pfd = ft.cell(0, 0).paragraphs[0]
    pfd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfd.add_run("_______________________________\n").bold = True
    rnd = pfd.add_run((meta.docente or "Docente de la Sesión") + "\n")
    rnd.bold = True
    rnd.font.size = Pt(9.5)
    rcd = pfd.add_run("Docente de la Sesión")
    rcd.font.size = Pt(8.5)

    pfs = ft.cell(0, 1).paragraphs[0]
    pfs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfs.add_run("_______________________________\n").bold = True
    rns = pfs.add_run((meta.director or "Director de la I.E.") + "\n")
    rns.bold = True
    rns.font.size = Pt(9.5)
    rcs = pfs.add_run("Director de la I.E.")
    rcs.font.size = Pt(8.5)

    # ── FICHA DE TRABAJO ──
    if doc_v1.fichaTrabajo and doc_v1.fichaTrabajo.titulo:
        doc.add_page_break()
        f_sec = doc.add_section()
        f_sec.top_margin = Inches(0.8)
        f_sec.bottom_margin = Inches(0.8)
        f_sec.left_margin = Inches(0.8)
        f_sec.right_margin = Inches(0.8)

        ft_tbl = doc.add_table(rows=1, cols=1)
        ft_tbl.autofit = False
        ft_tbl.rows[0].cells[0].width = Inches(6.67)
        add_table_borders_black(ft_tbl, sz='4')
        set_cell_text_white_bold(ft_tbl.cell(0, 0), "FICHA DE TRABAJO DE APRENDIZAJE INDEPENDIENTE", font_size_pt=11.5)
        set_cell_background(ft_tbl.cell(0, 0), '2980B9')
        set_cell_margins(ft_tbl.cell(0, 0), top=120, bottom=120, left=180, right=180)

        doc.add_paragraph().paragraph_format.space_before = Pt(10)

        stud_tbl = doc.add_table(rows=1, cols=2)
        stud_tbl.autofit = False
        add_table_borders(stud_tbl, color='CBD5E1', sz='4')

        set_cell_margins(stud_tbl.cell(0, 0), top=80, bottom=80, left=120, right=120)
        p_st1 = stud_tbl.cell(0, 0).paragraphs[0]
        r_st1 = p_st1.add_run("Estudiante: __________________________________________________")
        r_st1.bold = True
        r_st1.font.size = Pt(9.5)

        set_cell_margins(stud_tbl.cell(0, 1), top=80, bottom=80, left=120, right=120)
        p_st2 = stud_tbl.cell(0, 1).paragraphs[0]
        p_st2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_st2 = p_st2.add_run("Grado y Sección: ________________")
        r_st2.bold = True
        r_st2.font.size = Pt(9.5)

        for row in stud_tbl.rows:
            row.cells[0].width = Inches(4.5)
            row.cells[1].width = Inches(2.17)

        doc.add_paragraph().paragraph_format.space_before = Pt(8)

        p_ft_t = doc.add_paragraph()
        p_ft_t.paragraph_format.space_after = Pt(4)
        rf_t = p_ft_t.add_run("Actividad: " + doc_v1.fichaTrabajo.titulo.upper())
        rf_t.bold = True
        rf_t.font.size = Pt(12)
        rf_t.font.color.rgb = ACCENT_RGB

        if doc_v1.fichaTrabajo.indicaciones:
            p_ft_ind = doc.add_paragraph()
            p_ft_ind.paragraph_format.space_after = Pt(12)
            p_ft_ind.paragraph_format.line_spacing = 1.15
            p_ft_ind.add_run("Indicaciones: ").bold = True
            p_ft_ind.runs[0].font.size = Pt(9.5)
            p_ft_ind.add_run(doc_v1.fichaTrabajo.indicaciones).font.size = Pt(9.5)
            p_ft_ind.runs[1].font.italic = True

        act_html = doc_v1.fichaTrabajo.actividades or ""
        soup_act = BeautifulSoup(act_html, 'html.parser')

        def add_act_element(element, is_bold=False, is_italic=False, list_style=None):
            if isinstance(element, Tag):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(element.get_text().strip())
                    r.bold = True
                    r.font.size = Pt(11)
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    append_html_to_cell_or_paragraph(p, str(element), default_font_size=9.5)
                elif element.name == 'ul':
                    for li in element.find_all('li', recursive=False):
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.line_spacing = 1.15
                        rb = p.add_run(u"\u25cf ")
                        rb.font.size = Pt(8.5)
                        rb.font.color.rgb = PRIMARY_RGB
                        append_html_to_cell_or_paragraph(p, str(li), default_font_size=9.5)
                elif element.name == 'ol':
                    for idx_li, li in enumerate(element.find_all('li', recursive=False)):
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.line_spacing = 1.15
                        p.add_run(f"{idx_li+1}. ").bold = True
                        append_html_to_cell_or_paragraph(p, str(li), default_font_size=9.5)
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    if rows:
                        max_c = max(len(r.find_all(['td', 'th'])) for r in rows)
                        if max_c > 0:
                            tbl = doc.add_table(rows=len(rows), cols=max_c)
                            tbl.autofit = True
                            add_table_borders(tbl)
                            for ri, row in enumerate(rows):
                                cells = row.find_all(['td', 'th'])
                                for ci, cell_h in enumerate(cells):
                                    if ci < max_c:
                                        c = tbl.cell(ri, ci)
                                        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                                        p_cell = c.paragraphs[0]
                                        p_cell.paragraph_format.space_after = Pt(0)
                                        p_cell.paragraph_format.line_spacing = 1.1
                                        if cell_h.name == 'th':
                                            set_cell_background(c, 'F2F2F2')
                                            append_html_to_cell_or_paragraph(p_cell, f"<strong>{cell_h.get_text()}</strong>")
                                        else:
                                            append_html_to_cell_or_paragraph(p_cell, str(cell_h))
                else:
                    for child in element.children:
                        add_act_element(child, is_bold, is_italic)

        for child in soup_act.children:
            add_act_element(child)

    # ── LISTA DE COTEJO ──
    alumnos = doc_v1.listaCotejo.alumnos if doc_v1.listaCotejo else []
    criterios = doc_v1.listaCotejo.criterios if (doc_v1.listaCotejo and doc_v1.listaCotejo.criterios) else doc_v1.proposito.criterios
    if alumnos and len(alumnos) > 0 and len(criterios) > 0:
        doc.add_page_break()

        lc_sec = doc.add_section()
        lc_sec.page_width = Inches(11.69)
        lc_sec.page_height = Inches(8.27)
        lc_sec.top_margin = Inches(0.6)
        lc_sec.bottom_margin = Inches(0.6)
        lc_sec.left_margin = Inches(0.6)
        lc_sec.right_margin = Inches(0.6)

        lc_tbl = doc.add_table(rows=1, cols=1)
        lc_tbl.autofit = False
        lc_tbl.rows[0].cells[0].width = Inches(10.49)
        add_table_borders_black(lc_tbl)
        set_cell_text_white_bold(lc_tbl.cell(0, 0), "LISTA DE COTEJO DE EVALUACION FORMATIVA", font_size_pt=12)
        set_cell_background(lc_tbl.cell(0, 0), '2C3E50')
        set_cell_margins(lc_tbl.cell(0, 0), top=120, bottom=120, left=180, right=180)

        doc.add_paragraph().paragraph_format.space_before = Pt(8)

        lch_tbl = doc.add_table(rows=1, cols=4)
        lch_tbl.autofit = False
        add_table_borders_black(lch_tbl)

        set_cell_background(lch_tbl.cell(0, 0), PEACH)
        set_cell_margins(lch_tbl.cell(0, 0), top=80, bottom=80, left=120, right=120)
        lch_tbl.cell(0, 0).paragraphs[0].add_run("IE / Area").bold = True
        lch_tbl.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(8.5)

        set_cell_background(lch_tbl.cell(0, 1), GRAY_VAL)
        set_cell_margins(lch_tbl.cell(0, 1), top=80, bottom=80, left=120, right=120)
        lch_tbl.cell(0, 1).paragraphs[0].add_run(f"{meta.institucion or 'IE'} / {meta.area or 'Área'}")

        set_cell_background(lch_tbl.cell(0, 2), PEACH)
        set_cell_margins(lch_tbl.cell(0, 2), top=80, bottom=80, left=120, right=120)
        lch_tbl.cell(0, 2).paragraphs[0].add_run("Grado / Seccion").bold = True
        lch_tbl.cell(0, 2).paragraphs[0].runs[0].font.size = Pt(8.5)

        set_cell_background(lch_tbl.cell(0, 3), GRAY_VAL)
        set_cell_margins(lch_tbl.cell(0, 3), top=80, bottom=80, left=120, right=120)
        lch_tbl.cell(0, 3).paragraphs[0].add_run(f"{meta.grado or ''} \"{meta.seccion or ''}\"")

        for row in lch_tbl.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(4.5)
            row.cells[2].width = Inches(1.5)
            row.cells[3].width = Inches(2.99)

        doc.add_paragraph().paragraph_format.space_before = Pt(8)

        num_cols = 2 + len(criterios) * 2
        lct = doc.add_table(rows=2 + len(alumnos), cols=num_cols)
        lct.autofit = False
        add_table_borders_black(lct)

        lct.cell(0, 0).merge(lct.cell(1, 0))
        lct.cell(0, 1).merge(lct.cell(1, 1))

        for ci, crit in enumerate(criterios):
            sc = 2 + ci * 2
            lct.cell(0, sc).merge(lct.cell(0, sc + 1))
            lct.cell(0, sc).paragraphs[0].text = f"Criterio {ci + 1}: {crit}"
            lct.cell(1, sc).paragraphs[0].text = "SI"
            lct.cell(1, sc + 1).paragraphs[0].text = "NO"

        CRIT_COLORS = ['D9E1F2', 'FADBD8', 'D5F5E3', 'FCF3CF', 'FDE8D8', 'E8DAEF']
        SUBCRIT_COLORS = ['BDD7EE', 'FADBD8', 'A9DFBF', 'F9E79F', 'FAD7A0', 'D7BDE2']

        def _lcfmt(cell, width_in, font_size_pt, bold=False, ctr=False, bg=None):
            cell.width = Inches(width_in)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            if bg:
                set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if ctr:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = bold
                run.font.size = Pt(font_size_pt)

        _lcfmt(lct.cell(0, 0), 0.35, 8.5, bold=True, ctr=True, bg='FFF2CC')
        _lcfmt(lct.cell(0, 1), 2.2, 8.5, bold=True, ctr=False, bg='FFF2CC')
        for ci, crit in enumerate(criterios):
            sc = 2 + ci * 2
            _lcfmt(lct.cell(0, sc), 0.7, 7.5, bold=True, ctr=True, bg=CRIT_COLORS[ci % len(CRIT_COLORS)])
            _lcfmt(lct.cell(1, sc), 0.35, 8, bold=True, ctr=True, bg=SUBCRIT_COLORS[ci % len(SUBCRIT_COLORS)])
            _lcfmt(lct.cell(1, sc + 1), 0.35, 8, bold=True, ctr=True, bg=SUBCRIT_COLORS[ci % len(SUBCRIT_COLORS)])
        for ri, stud in enumerate(alumnos):
            rn = 2 + ri
            lct.cell(rn, 0).text = str(ri + 1)
            lct.cell(rn, 1).text = "" if stud.startswith("Estudiante ") else stud
            _lcfmt(lct.cell(rn, 0), 0.35, 8.5, bold=True, ctr=True)
            _lcfmt(lct.cell(rn, 1), 2.2, 8.5, bold=False, ctr=False)
            for ci in range(len(criterios)):
                sc = 2 + ci * 2
                _lcfmt(lct.cell(rn, sc), 0.35, 8, ctr=True)
                _lcfmt(lct.cell(rn, sc + 1), 0.35, 8, ctr=True)

    _apply_presentation(doc, doc_v1)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
