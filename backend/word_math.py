"""Render híbrido de LaTeX para documentos Word.

Ruta de salida:
1. OMML nativo y editable para la notación educativa habitual.
2. PNG transparente generado por Matplotlib para expresiones no soportadas.
3. Texto LaTeX original si ninguna de las rutas anteriores está disponible.

El módulo no modifica el HTML ni el resto de los runs del párrafo.
"""
from __future__ import annotations

import io
import re
from functools import lru_cache
from typing import Iterable

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_LATEX_SEGMENTS = re.compile(
    r"(\$\$[\s\S]+?\$\$|\$(?:\\.|[^$\\])+?\$|\\\[[\s\S]+?\\\]|\\\((?:\\.|[^\\])+?\\\))"
)


class UnsupportedLatex(ValueError):
    """La expresión debe pasar al render de imagen."""


def split_latex_segments(text: str):
    """Devuelve tuplas (es_formula, contenido, es_bloque, fuente_original)."""
    cursor = 0
    for match in _LATEX_SEGMENTS.finditer(text or ""):
        if match.start() > cursor:
            raw = text[cursor:match.start()]
            yield False, raw, False, raw
        raw = match.group(0)
        display = raw.startswith("$$") or raw.startswith(r"\[")
        if raw.startswith("$$"):
            latex = raw[2:-2]
        else:
            latex = raw[2:-2] if raw.startswith((r"\[", r"\(")) else raw[1:-1]
        yield True, latex.strip(), display, raw
        cursor = match.end()
    if cursor < len(text or ""):
        raw = text[cursor:]
        yield False, raw, False, raw


def contains_latex(text: str) -> bool:
    return bool(_LATEX_SEGMENTS.search(text or ""))


def _math_text(value: str):
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    if value[:1].isspace() or value[-1:].isspace():
        text.set(_XML_SPACE, "preserve")
    text.text = value
    run.append(text)
    return run


def _container(name: str, children: Iterable):
    element = OxmlElement(name)
    for child in children:
        element.append(child)
    return element


def _script(name: str, base: list, value: list):
    root = OxmlElement(name)
    root.append(_container("m:e", base))
    slot = "m:sub" if name == "m:sSub" else "m:sup"
    root.append(_container(slot, value))
    return root


def _sub_sup(base: list, sub: list, sup: list):
    root = OxmlElement("m:sSubSup")
    root.append(_container("m:e", base))
    root.append(_container("m:sub", sub))
    root.append(_container("m:sup", sup))
    return root


_COMMANDS = {
    "pm": "±", "mp": "∓", "times": "×", "div": "÷", "cdot": "·",
    "le": "≤", "leq": "≤", "ge": "≥", "geq": "≥", "neq": "≠",
    "approx": "≈", "equiv": "≡", "in": "∈", "notin": "∉",
    "rightarrow": "→", "leftarrow": "←", "leftrightarrow": "↔",
    "Rightarrow": "⇒", "Leftarrow": "⇐", "Leftrightarrow": "⇔",
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ε", "theta": "θ", "lambda": "λ", "mu": "μ",
    "pi": "π", "rho": "ρ", "sigma": "σ", "phi": "φ", "omega": "ω",
    "Delta": "Δ", "Sigma": "Σ", "Phi": "Φ", "Omega": "Ω",
    "infty": "∞", "partial": "∂", "sum": "∑", "prod": "∏", "int": "∫",
    "forall": "∀", "exists": "∃", "ldots": "…", "cdots": "⋯",
}


class _Parser:
    def __init__(self, latex: str):
        self.value = latex
        self.pos = 0

    def parse(self, stop: str | None = None) -> list:
        output = []
        plain = []

        def flush():
            if plain:
                output.append(_math_text("".join(plain)))
                plain.clear()

        while self.pos < len(self.value):
            char = self.value[self.pos]
            if stop and char == stop:
                break
            if char.isspace():
                plain.append(" ")
                self.pos += 1
                continue
            if char not in "\\{}_^":
                if self.pos + 1 >= len(self.value) or self.value[self.pos + 1] not in "_^":
                    plain.append(char)
                    self.pos += 1
                    continue
                flush()
                self.pos += 1
                atom = [_math_text(char)]
            else:
                flush()
                if char == "}":
                    if stop:
                        break
                    raise UnsupportedLatex("Llave de cierre inesperada")
                atom = self._atom()

            sub = sup = None
            while self.pos < len(self.value) and self.value[self.pos] in "_^":
                marker = self.value[self.pos]
                self.pos += 1
                script_value = self._argument()
                if marker == "_":
                    sub = script_value
                else:
                    sup = script_value
            if sub is not None and sup is not None:
                atom = [_sub_sup(atom, sub, sup)]
            elif sub is not None:
                atom = [_script("m:sSub", atom, sub)]
            elif sup is not None:
                atom = [_script("m:sSup", atom, sup)]
            output.extend(atom)
        flush()
        return output

    def _atom(self) -> list:
        if self.pos >= len(self.value):
            return []
        char = self.value[self.pos]
        if char == "{":
            return self._group()
        if char == "\\":
            return self._command()
        if char in "_^":
            raise UnsupportedLatex("Subíndice o superíndice sin base")
        self.pos += 1
        return [_math_text(char)]

    def _group(self) -> list:
        if self.value[self.pos] != "{":
            raise UnsupportedLatex("Se esperaba un grupo")
        self.pos += 1
        content = self.parse(stop="}")
        if self.pos >= len(self.value) or self.value[self.pos] != "}":
            raise UnsupportedLatex("Grupo LaTeX sin cerrar")
        self.pos += 1
        return content

    def _argument(self) -> list:
        if self.pos >= len(self.value):
            raise UnsupportedLatex("Argumento LaTeX ausente")
        return self._group() if self.value[self.pos] == "{" else self._atom()

    def _command(self) -> list:
        self.pos += 1
        start = self.pos
        while self.pos < len(self.value) and self.value[self.pos].isalpha():
            self.pos += 1
        command = self.value[start:self.pos]
        if not command and self.pos < len(self.value):
            escaped = self.value[self.pos]
            self.pos += 1
            if escaped in "{}_$%#&":
                return [_math_text(escaped)]
            if escaped in ",;:! ":
                return [_math_text(" ")]
            raise UnsupportedLatex(f"Comando no soportado: \\{escaped}")

        if command == "frac":
            numerator = self._argument()
            denominator = self._argument()
            fraction = OxmlElement("m:f")
            fraction.append(_container("m:num", numerator))
            fraction.append(_container("m:den", denominator))
            return [fraction]
        if command == "sqrt":
            radicand = self._argument()
            radical = OxmlElement("m:rad")
            props = OxmlElement("m:radPr")
            hide = OxmlElement("m:degHide")
            hide.set(qn("m:val"), "1")
            props.append(hide)
            radical.append(props)
            radical.append(_container("m:e", radicand))
            return [radical]
        if command in {"text", "mathrm", "mathbf", "mathit"}:
            return self._argument()
        if command in {"left", "right"}:
            return self._atom()
        if command in _COMMANDS:
            return [_math_text(_COMMANDS[command])]
        raise UnsupportedLatex(f"Comando no soportado: \\{command}")


def latex_to_omml(latex: str):
    """Convierte el subconjunto seguro de LaTeX a un nodo m:oMath editable."""
    if not latex or len(latex) > 4000 or r"\begin" in latex or r"\end" in latex:
        raise UnsupportedLatex("Expresión vacía, demasiado grande o con entorno complejo")
    parser = _Parser(latex)
    children = parser.parse()
    if parser.pos != len(parser.value):
        raise UnsupportedLatex("No se pudo consumir toda la expresión")
    equation = OxmlElement("m:oMath")
    for child in children:
        equation.append(child)
    return equation


@lru_cache(maxsize=128)
def _render_formula_png(latex: str, font_size: float) -> bytes | None:
    """Renderiza MathText bajo demanda. La dependencia es opcional en runtime."""
    try:
        from matplotlib.font_manager import FontProperties
        from matplotlib.mathtext import math_to_image

        stream = io.BytesIO()
        math_to_image(
            f"${latex}$",
            stream,
            prop=FontProperties(size=max(9.0, float(font_size) * 1.35)),
            dpi=200,
            format="png",
            color="black",
        )
        # math_to_image usa fondo blanco. Convertirlo a alfa real evita que la
        # fórmula se vea como un rectángulo sobre celdas coloreadas de Word.
        from PIL import Image

        stream.seek(0)
        with Image.open(stream).convert("RGBA") as source:
            transparent = Image.new("RGBA", source.size)
            pixels = []
            for red, green, blue, alpha in source.getdata():
                luminance = round(0.299 * red + 0.587 * green + 0.114 * blue)
                ink_alpha = min(alpha, 255 - luminance)
                pixels.append((0, 0, 0, ink_alpha))
            transparent.putdata(pixels)
            output = io.BytesIO()
            transparent.save(output, format="PNG", dpi=(200, 200))
        return output.getvalue()
    except Exception:
        return None


def _append_formula_image(paragraph, latex: str, font_size: float, display: bool) -> bool:
    image = _render_formula_png(latex, font_size)
    if not image:
        return False
    try:
        from PIL import Image

        stream = io.BytesIO(image)
        with Image.open(stream) as bitmap:
            width_px, height_px = bitmap.size
        stream.seek(0)
        width_inches = max(0.12, width_px / 200)
        height_inches = max(0.12, height_px / 200)
        max_width = 6.1 if display else 2.8
        scale = min(1.0, max_width / width_inches)
        if not display and height_inches * scale < 0.16:
            scale = min(max_width / width_inches, 0.16 / height_inches)
        paragraph.add_run().add_picture(stream, width=Inches(width_inches * scale))
        return True
    except Exception:
        return False


def append_mixed_text(
    paragraph,
    text: str,
    *,
    font_name: str = "Arial",
    font_size: float = 9.5,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    font_color=None,
) -> list[str]:
    """Añade texto y fórmulas al mismo párrafo sin aplanar sus estilos.

    Retorna las rutas usadas (``text``, ``omml``, ``image`` o ``latex-text``),
    lo que permite probar el fallback sin inspeccionar detalles internos.
    """
    routes = []
    segments = list(split_latex_segments(text or ""))
    only_display = len(segments) == 1 and segments[0][0] and segments[0][2]
    if only_display:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for is_formula, value, display, original in segments:
        if not is_formula:
            if not value:
                continue
            run = paragraph.add_run(value)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if font_color is not None:
                run.font.color.rgb = font_color
            routes.append("text")
            continue
        try:
            paragraph._p.append(latex_to_omml(value))
            routes.append("omml")
        except UnsupportedLatex:
            if _append_formula_image(paragraph, value, font_size, display):
                routes.append("image")
            else:
                run = paragraph.add_run(original)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                if font_color is not None:
                    run.font.color.rgb = font_color
                routes.append("latex-text")
    return routes
